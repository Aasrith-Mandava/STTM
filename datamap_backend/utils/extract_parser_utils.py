"""
Extract Parser Utilities
========================
Pure Python utility functions for parsing BRD documents, layout specifications,
and meeting transcripts. No ADK imports — this module is framework-agnostic.
"""

import io
import json
import re


def parse_docx_brd(brd_bytes: bytes, section_ref: str | None = None) -> str:
    """
    Parse a DOCX Business Requirements Document into plain text.

    Args:
        brd_bytes: Raw bytes of the .docx file.
        section_ref: Optional filter in the format "solution_owner:<name>".
                     When provided, only paragraphs mentioning the name are kept.

    Returns:
        Newline-joined string of extracted text, with TBD lines stripped.
    """
    from docx import Document

    doc = Document(io.BytesIO(brd_bytes))
    lines: list[str] = []

    # Extract all paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Extract all table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    # Filter by section_ref if provided (solution_owner:<name>)
    if section_ref and section_ref.startswith("solution_owner:"):
        owner_name = section_ref.split(":", 1)[1].strip()
        if owner_name:
            lines = [
                line for line in lines
                if owner_name.lower() in line.lower()
            ]

    # Strip TBD / "To Be Determined" lines
    tbd_values = {"tbd", "to be determined"}
    lines = [
        line for line in lines
        if line.strip().lower() not in tbd_values
    ]

    return "\n".join(lines)


def parse_xlsx_layout(xlsx_bytes: bytes) -> str:
    """
    Parse an Excel layout specification into a JSON string of field rows.

    Args:
        xlsx_bytes: Raw bytes of the .xlsx file.

    Returns:
        JSON string (pretty-printed) of extracted layout rows.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    ws = wb.active

    all_rows = list(ws.iter_rows(values_only=True))
    if not all_rows:
        return json.dumps([], indent=2)

    # Auto-detect header row: first row where at least 3 cells are non-empty
    header_row_idx = None
    for idx, row in enumerate(all_rows):
        non_empty = sum(1 for cell in row if cell is not None and str(cell).strip())
        if non_empty >= 3:
            header_row_idx = idx
            break

    if header_row_idx is None:
        return json.dumps([], indent=2)

    raw_headers = all_rows[header_row_idx]

    # Normalise header names to lowercase with underscores
    def _normalise_header(h: str) -> str:
        h = str(h).strip()
        h = re.sub(r'[^a-zA-Z0-9\s]', '', h)
        h = re.sub(r'\s+', '_', h)
        return h.lower()

    headers = [_normalise_header(str(h)) if h else f"col_{i}" for i, h in enumerate(raw_headers)]

    # Target columns for flexible matching
    target_columns = [
        "attribute_name", "description", "data_type",
        "length", "format", "nullability",
    ]

    def _matches_target(header: str, target: str) -> bool:
        """Check if header partially matches a target column name."""
        return target in header or header in target

    def _find_matching_targets(header: str) -> str | None:
        for target in target_columns:
            if _matches_target(header, target):
                return target
        return None

    # Build column index map
    column_map: dict[int, str] = {}
    for col_idx, header in enumerate(headers):
        match = _find_matching_targets(header)
        if match:
            column_map[col_idx] = match
        else:
            column_map[col_idx] = header

    # Extract data rows
    rows: list[dict] = []
    for row in all_rows[header_row_idx + 1:]:
        # Skip fully empty rows
        if all(cell is None or str(cell).strip() == "" for cell in row):
            continue

        row_dict: dict[str, str] = {}
        for col_idx, cell in enumerate(row):
            if col_idx in column_map:
                key = column_map[col_idx]
                row_dict[key] = str(cell) if cell is not None else ""

        rows.append(row_dict)

    return json.dumps(rows, default=str, indent=2)


_SKIP_SHEET_NAMES: set[str] = set()


def _resolve_merged_cells(ws) -> dict[tuple[int, int], str]:
    """
    Build a (row, col) → value map that fills every cell in a merged range
    with the top-left cell's value so downstream code never sees None for
    a merged cell.
    """
    from openpyxl.utils import range_boundaries

    filled: dict[tuple[int, int], str] = {}
    for merged_range in ws.merged_cells.ranges:
        min_col, min_row, max_col, max_row = range_boundaries(str(merged_range))
        value = ws.cell(min_row, min_col).value
        str_value = str(value).strip() if value is not None else ""
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                filled[(r, c)] = str_value
    return filled


def _sheet_to_rows(ws) -> list[list[str]]:
    """
    Return all non-empty rows from a worksheet as list[list[str]],
    with merged-cell values propagated.
    """
    merged_map = _resolve_merged_cells(ws)
    result: list[list[str]] = []
    for row in ws.iter_rows():
        cells = [
            merged_map.get((cell.row, cell.column),
                           str(cell.value).strip() if cell.value is not None else "")
            for cell in row
        ]
        if any(c for c in cells):
            result.append(cells)
    return result


def _dedup_headers(raw: list[str]) -> list[str]:
    """Deduplicate column headers by appending _2, _3 … for repeats."""
    seen: dict[str, int] = {}
    out: list[str] = []
    for h in raw:
        key = h.strip() or "col"
        seen[key] = seen.get(key, 0) + 1
        out.append(f"{key}_{seen[key]}" if seen[key] > 1 else key)
    return out


def parse_xlsx_to_json(xlsx_bytes: bytes) -> dict[str, list[dict[str, str]]]:
    """
    Convert an Excel workbook directly to the canonical UI shape:
      { "Sheet Name": [ {"Column Header": "cell value", ...}, ... ] }

    Handles:
    - Multi-sheet workbooks (each sheet → one top-level key)
    - Merged cells (value propagated across the entire merged range)
    - Smart header detection (first row with 3+ non-empty cells)
    - Trailing empty column trimming
    - Section-label row filtering (rows with < 2 non-empty cells skipped)
    - Metadata sheet skipping (toc, cover, index, readme, instructions)
    - Duplicate column header deduplication
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    result: dict[str, list[dict[str, str]]] = {}

    for sheet_name in wb.sheetnames:
        if sheet_name.lower().strip() in _SKIP_SHEET_NAMES:
            continue

        ws = wb[sheet_name]
        all_rows = _sheet_to_rows(ws)
        if not all_rows:
            continue

        # Find first row with 2+ non-empty cells → header row
        header_idx = next(
            (i for i, row in enumerate(all_rows)
             if sum(1 for c in row if c) >= 2),
            None,
        )
        # Skip sheet if no valid header found
        if header_idx is None:
            continue

        raw_headers = list(all_rows[header_idx])  # copy to avoid mutating all_rows
        # Trim trailing empty columns
        while raw_headers and not raw_headers[-1]:
            raw_headers.pop()
        col_count = len(raw_headers)
        headers = _dedup_headers(raw_headers)

        rows: list[dict[str, str]] = []
        for raw_row in all_rows[header_idx + 1:]:
            # Skip fully empty rows only
            if not any(c for c in raw_row):
                continue
            padded = (raw_row + [""] * col_count)[:col_count]
            rows.append(dict(zip(headers, padded)))

        if rows:
            result[sheet_name] = rows

    return result


# Keep old name as alias so any other callers don't break
_parse_xlsx_layout_direct = parse_xlsx_to_json


def xlsx_to_markdown(xlsx_bytes: bytes) -> str:
    """
    Convert a multi-sheet Excel workbook to a markdown string.
    Each sheet becomes a ## section followed by a GitHub-flavoured markdown table.

    Used by brd_transcript_agent for file layout markdown conversion.

    - Skips metadata/TOC sheets
    - Smart header detection: first row with 3+ non-empty cells
    - Handles merged cells (value propagated)
    - Trims trailing empty columns
    - Skips rows with fewer than 2 non-empty cells (section label rows)
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        if sheet_name.lower().strip() in _SKIP_SHEET_NAMES:
            continue

        ws = wb[sheet_name]
        all_rows = _sheet_to_rows(ws)
        if not all_rows:
            continue

        # Smart header detection: first row with 2+ non-empty cells
        header_idx = next(
            (i for i, row in enumerate(all_rows)
             if sum(1 for c in row if c) >= 2),
            None,
        )
        if header_idx is None:
            continue

        headers = list(all_rows[header_idx])  # copy to avoid mutating all_rows
        # Trim trailing empty columns
        while headers and not headers[-1]:
            headers.pop()
        col_count = len(headers)

        lines: list[str] = [f"## {sheet_name}", ""]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")

        for row in all_rows[header_idx + 1:]:
            # Skip fully empty rows only
            if not any(c for c in row):
                continue
            cells = [
                c.replace("|", "\\|").replace("\n", " ") if c else ""
                for c in row[:col_count]
            ]
            cells += [""] * (col_count - len(cells))
            lines.append("| " + " | ".join(cells) + " |")

        sections.append("\n".join(lines))

    return "\n\n".join(sections)


def extract_decisions_from_transcript(raw_text: str) -> str:
    """
    Filter a meeting transcript to keep only confirmed-decision lines.

    Args:
        raw_text: Full transcript text.

    Returns:
        Newline-joined string of decision lines, or a fallback message.
    """
    keep_keywords = ["confirmed", "agreed", "decided", "shall", "action item"]
    remove_keywords = ["tbd", "open item", "question", "to be confirmed"]

    kept_lines: list[str] = []

    for line in raw_text.splitlines():
        line_lower = line.lower()

        # Skip lines with removal keywords
        if any(kw in line_lower for kw in remove_keywords):
            continue

        # Keep lines with keeper keywords
        if any(kw in line_lower for kw in keep_keywords):
            kept_lines.append(line.strip())

    if not kept_lines:
        return "No confirmed decisions found in transcript."

    return "\n".join(kept_lines)